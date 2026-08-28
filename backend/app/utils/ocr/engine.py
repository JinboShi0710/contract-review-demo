# -*- coding: utf-8 -*-
"""
PaddleOCR API 封装
使用远程 API 服务
"""
import json
import time
from pathlib import Path
from typing import Tuple, List, Optional
import requests
from pypdf import PdfReader
from docx import Document
from app.core.config import settings


class OCREngine:
    """PaddleOCR API 引擎封装"""

    _instance: Optional["OCREngine"] = None

    API_URL = "https://paddleocr.aistudio-app.com/api/v2/ocr/jobs"
    TOKEN = settings.OCR_API_TOKEN
    MODEL = "PaddleOCR-VL-1.5"

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        if self._initialized:
            return
        print("PaddleOCR API 引擎初始化完成")
        self._initialized = True

    def recognize(self, file_path: str) -> Tuple[List[dict], Optional[float], int]:
        """
        识别图片中的文字，直接返回原始文本

        Returns:
            (
                results: List[{"text": str, "confidence": float, "bbox": ...}],
                avg_confidence: float,
                page_count: int
            )
        """
        suffix = Path(file_path).suffix.lower()

        # Word 文档直接读取段落和表格。DOCX 本身不保存可靠的最终分页结果，
        # 因此这里作为一个逻辑文档处理，不虚构页码。
        if suffix == ".docx":
            document = Document(file_path)
            blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        blocks.append("\t".join(cells))
            text = "\n".join(blocks).strip()
            if not text:
                raise RuntimeError("DOCX 中未提取到可审核文字")
            print(f"DOCX 直接提取完成: {len(text)} 字符", flush=True)
            return [{"text": text, "page": 1, "confidence": None}], None, 1

        # TXT/Markdown 直接读取，兼容常见中文编码。
        if suffix in (".txt", ".md"):
            raw = Path(file_path).read_bytes()
            text = None
            for encoding in ("utf-8-sig", "gb18030", "utf-16"):
                try:
                    text = raw.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if text is None or not text.strip():
                raise RuntimeError("文本文件为空或编码无法识别")
            text = text.strip()
            print(f"文本文件直接提取完成: {len(text)} 字符", flush=True)
            return [{"text": text, "page": 1, "confidence": None}], None, 1

        # 文本型 PDF 直接读取文字：更快、页数准确，也避免把敏感合同
        # 无必要地上传到第三方 OCR。仅扫描件和图片走远程 OCR。
        if suffix == ".pdf":
            try:
                reader = PdfReader(file_path)
                page_items = []
                total_chars = 0
                for page_no, page in enumerate(reader.pages, start=1):
                    text = (page.extract_text() or "").strip()
                    total_chars += len(text)
                    if text:
                        page_items.append({
                            "text": text,
                            "page": page_no,
                            "confidence": None,
                        })

                # 平均每页至少 30 个可提取字符时，认为是文本型 PDF。
                # 否则回退到远程 OCR，以兼容扫描件或只有少量页眉文字的文件。
                if reader.pages and total_chars >= len(reader.pages) * 30:
                    print(f"PDF 直接提取完成: {len(reader.pages)} 页, {total_chars} 字符", flush=True)
                    return page_items, None, len(reader.pages)
                print("PDF 可提取文本过少，改用 PaddleOCR", flush=True)
            except Exception as exc:
                print(f"PDF 直接提取失败，改用 PaddleOCR: {exc}", flush=True)

        headers = {"Authorization": f"bearer {self.TOKEN}"}

        optional_payload = {
            "useDocOrientationClassify": False,
            "useDocUnwarping": False,
            "useChartRecognition": False,
        }

        data = {
            "model": self.MODEL,
            "optionalPayload": json.dumps(optional_payload)
        }

        with open(file_path, "rb") as f:
            files = {"file": f}
            job_response = requests.post(
                self.API_URL,
                headers=headers,
                data=data,
                files=files,
                timeout=120
            )

        if job_response.status_code != 200:
            raise RuntimeError(f"OCR API 请求失败: {job_response.status_code} - {job_response.text}")

        job_id = job_response.json()["data"]["jobId"]

        # 轮询结果
        while True:
            job_result = requests.get(f"{self.API_URL}/{job_id}", headers=headers, timeout=30)
            if job_result.status_code != 200:
                raise RuntimeError(f"OCR 结果查询失败: {job_result.status_code}")

            state = job_result.json()["data"]["state"]

            if state == "done":
                result_url = job_result.json()["data"]["resultUrl"]["jsonUrl"]
                jsonl_response = requests.get(result_url, timeout=30)
                jsonl_response.raise_for_status()

                # 解析每页结果，保留页码信息
                all_items = []
                page_count = 0
                for line in jsonl_response.text.strip().split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    result_data = json.loads(line)["result"]
                    layout_results = result_data.get("layoutParsingResults", [])
                    for layout_result in layout_results:
                        page_count += 1
                        markdown_text = layout_result.get("markdown", {}).get("text", "")
                        if markdown_text:
                            all_items.append({
                                "text": markdown_text,
                                "page": page_count,
                                "confidence": None,
                            })

                    # 兼容某些 API 返回一行一个页面、但没有 layoutParsingResults 的情况。
                    if not layout_results:
                        page_count += 1

                return all_items, None, page_count

            elif state == "failed":
                error_msg = job_result.json()["data"].get("errorMsg", "未知错误")
                raise RuntimeError(f"OCR 处理失败: {error_msg}")

            elif state in ("pending", "running"):
                print(f"OCR 处理中... state: {state}")
                time.sleep(5)
            else:
                raise RuntimeError(f"OCR 未知状态: {state}")

    def is_original(self, file_path: str) -> Tuple[bool, float]:
        """
        判断文档是原件还是复印件
        """
        from PIL import Image, ImageFilter
        import numpy as np

        try:
            img = Image.open(file_path).convert("L")
            img_array = np.array(img)

            laplacian = img.filter(ImageFilter.Kernel(
                size=(3, 3),
                kernel=[-1, -1, -1, -1, 8, -1, -1, -1, -1],
            ))
            laplacian_array = np.array(laplacian)
            sharpness = np.var(laplacian_array)

            is_original = sharpness > 1000

            return is_original, float(sharpness)
        except Exception as e:
            print(f"原件/复印件判断失败: {e}")
            return True, 0.0


# 单例
ocr_engine = OCREngine()
