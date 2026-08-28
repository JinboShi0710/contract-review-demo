# -*- coding: utf-8 -*-
"""
审核报告生成服务 — 支持 PDF 和 Word 双格式
"""
import io
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings
from app.models.contract import Contract, ContractRisk, ContractReport


# 注册中文字体（SimHei 黑体，Windows 内置）
_FONT_REGISTERED = False

def _ensure_chinese_font():
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return
    font_path = Path("C:/Windows/Fonts/simhei.ttf")
    if font_path.exists():
        pdfmetrics.registerFont(TTFont("SimHei", str(font_path)))
        _FONT_REGISTERED = True


# 规则类型分组顺序与标签
RULE_TYPE_ORDER = [
    "keyword",
    "regex",
    "blacklist",
    "llm_risk",
    "llm_compliance",
    "llm_completeness",
]

RULE_TYPE_LABELS = {
    "keyword": "关键词检测规则",
    "regex": "格式校验规则",
    "blacklist": "黑名单规则",
    "llm_risk": "LLM 风险识别规则",
    "llm_compliance": "LLM 合规检查规则",
    "llm_completeness": "LLM 完整性检查规则",
}

RISK_LEVEL_LABELS = {"high": "高", "medium": "中", "low": "低"}
MATCHED_LABELS = {0: "未执行", 1: "通过", 2: "风险"}
MAX_MATCHED_ITEMS = 5


def _build_conclusion(risks: List[ContractRisk]) -> tuple[str, dict]:
    """
    根据风险列表自动生成结论和计数。
    返回 (结论文字, {high, medium, low, total_risk})
    """
    counts = {"high": 0, "medium": 0, "low": 0}
    has_risk = False
    for risk in risks:
        if risk.matched == 2:
            has_risk = True
            level = risk.risk_level if risk.risk_level in counts else "low"
            counts[level] += 1

    total = counts["high"] + counts["medium"] + counts["low"]
    if not has_risk:
        conclusion = "审核通过，未发现风险条目。"
    else:
        conclusion = (
            f"存在风险，共发现 {total} 条风险（"
            f"高风险 {counts['high']} 条，"
            f"中风险 {counts['medium']} 条，"
            f"低风险 {counts['low']} 条）。"
        )
    return conclusion, counts


def _group_risks(risks: List[ContractRisk]) -> List[tuple[str, List[ContractRisk]]]:
    """按规则类型分组，按 RULE_TYPE_ORDER 排序，未知类型追加到末尾。"""
    groups: dict[str, List[ContractRisk]] = {}
    for risk in risks:
        key = risk.rule_type or "unknown"
        groups.setdefault(key, []).append(risk)

    ordered = []
    for rtype in RULE_TYPE_ORDER:
        if rtype in groups:
            ordered.append((rtype, groups.pop(rtype)))
    # 未知类型追加
    for rtype, items in groups.items():
        ordered.append((rtype, items))
    return ordered


def _truncate_matched_items(items) -> tuple[list, int]:
    """截断 matched_items，返回 (前5条, 总数)。"""
    if not items:
        return [], 0
    total = len(items)
    return list(items[:MAX_MATCHED_ITEMS]), total


class ReportGenerator:
    """审核报告生成器，支持 PDF 和 Word。"""

    def __init__(self):
        self.export_dir = settings.EXPORT_DIR
        self.export_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # 公共入口
    # ------------------------------------------------------------------

    def generate(
        self,
        contract: Contract,
        risks: List[ContractRisk],
        fmt: str = "pdf",
    ) -> ContractReport:
        """
        生成报告文件并返回 ContractReport 对象（未持久化，由调用方保存）。

        Args:
            contract: 合同对象
            risks: 风险列表
            fmt: 'pdf' 或 'word'

        Returns:
            ContractReport（report_path 已写入文件系统）
        """
        report_id = str(uuid.uuid4())
        conclusion, _ = _build_conclusion(risks)

        if fmt == "word":
            filename = f"report_{report_id}.docx"
            report_path = self.export_dir / filename
            self._build_word(contract, risks, conclusion, report_path)
            report_type = "word"
        else:
            filename = f"report_{report_id}.pdf"
            report_path = self.export_dir / filename
            self._build_pdf(contract, risks, conclusion, report_path)
            report_type = "pdf"

        return ContractReport(
            id=report_id,
            contract_id=contract.id,
            report_path=str(report_path),
            report_type=report_type,
            conclusion=conclusion,
        )

    # ------------------------------------------------------------------
    # PDF 生成
    # ------------------------------------------------------------------

    def _build_pdf(
        self,
        contract: Contract,
        risks: List[ContractRisk],
        conclusion: str,
        report_path: Path,
    ):
        _ensure_chinese_font()
        font_name = "SimHei" if _FONT_REGISTERED else "Helvetica"

        doc = SimpleDocTemplate(
            str(report_path),
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        story = []

        # 统一中文样式
        h1 = ParagraphStyle("CH1", fontName=font_name, fontSize=18, spaceAfter=20, alignment=1)
        h2 = ParagraphStyle("CH2", fontName=font_name, fontSize=14, spaceBefore=12, spaceAfter=6)
        h3 = ParagraphStyle("CH3", fontName=font_name, fontSize=12, spaceBefore=8, spaceAfter=4)
        normal = ParagraphStyle("CHN", fontName=font_name, fontSize=10, spaceAfter=4)
        cell_style = ParagraphStyle("CHC", fontName=font_name, fontSize=9)

        story.append(Paragraph("合同审核报告", h1))
        story.append(Spacer(1, 0.3 * cm))

        # --- 1. 合同基本信息 ---
        story.append(Paragraph("一、合同基本信息", h2))
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        info_data = [
            [Paragraph("合同名称", cell_style), Paragraph(contract.file_name, cell_style)],
            [Paragraph("合同编号", cell_style), Paragraph(contract.contract_type or "—", cell_style)],
            [Paragraph("上传时间", cell_style), Paragraph(contract.created_at.strftime("%Y-%m-%d %H:%M:%S"), cell_style)],
            [Paragraph("报告生成时间", cell_style), Paragraph(now_str, cell_style)],
        ]
        info_table = Table(info_data, colWidths=[4 * cm, 12 * cm])
        info_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F5F5F5")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.5 * cm))

        # --- 2. 审核结论 ---
        story.append(Paragraph("二、审核结论", h2))
        story.append(Paragraph(conclusion, normal))
        story.append(Spacer(1, 0.5 * cm))

        # --- 3. 风险详情 ---
        story.append(Paragraph("三、风险详情列表", h2))

        if not risks:
            story.append(Paragraph("暂无风险记录", normal))
        else:
            grouped = _group_risks(risks)
            for rule_type, items in grouped:
                label = RULE_TYPE_LABELS.get(rule_type, rule_type)
                story.append(Paragraph(f"▌ {label}", h3))

                for risk in items:
                    matched_label = MATCHED_LABELS.get(risk.matched or 0, "未知")
                    level_label = RISK_LEVEL_LABELS.get(risk.risk_level, risk.risk_level)

                    def p(text): return Paragraph(str(text), cell_style)

                    row_data = [
                        [p("规则名称"), p(risk.rule_name or risk.risk_type)],
                        [p("严重程度"), p(level_label)],
                        [p("执行结果"), p(matched_label)],
                    ]

                    # 命中内容（截断）
                    if risk.matched == 2 and risk.matched_items:
                        display_items, total = _truncate_matched_items(risk.matched_items)
                        items_text = "、".join(str(i) for i in display_items)
                        if total > MAX_MATCHED_ITEMS:
                            items_text += f"（共 {total} 条命中）"
                        row_data.append([p("命中内容"), p(items_text)])

                    # LLM 分析结果（有值才写）
                    if risk.execution_result:
                        row_data.append([p("分析说明"), p(risk.execution_result)])

                    rule_table = Table(row_data, colWidths=[3 * cm, 13 * cm])
                    rule_table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#FAFAFA")),
                        ("GRID", (0, 0), (-1, -1), 0.3, colors.lightgrey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]))
                    story.append(rule_table)
                    story.append(Spacer(1, 0.2 * cm))

                story.append(Spacer(1, 0.3 * cm))

        doc.build(story)

    # ------------------------------------------------------------------
    # Word 生成
    # ------------------------------------------------------------------

    def _build_word(
        self,
        contract: Contract,
        risks: List[ContractRisk],
        conclusion: str,
        report_path: Path,
    ):
        doc = Document()

        # 标题
        title = doc.add_heading("合同审核报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 1. 合同基本信息 ---
        doc.add_heading("一、合同基本信息", level=1)
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        info_rows = [
            ("合同名称", contract.file_name),
            ("合同编号", contract.contract_type or "—"),
            ("上传时间", contract.created_at.strftime("%Y-%m-%d %H:%M:%S")),
            ("报告生成时间", now_str),
        ]
        tbl = doc.add_table(rows=len(info_rows), cols=2)
        tbl.style = "Table Grid"
        for i, (key, val) in enumerate(info_rows):
            tbl.cell(i, 0).text = key
            tbl.cell(i, 1).text = val

        doc.add_paragraph()

        # --- 2. 审核结论 ---
        doc.add_heading("二、审核结论", level=1)
        doc.add_paragraph(conclusion)
        doc.add_paragraph()

        # --- 3. 风险详情 ---
        doc.add_heading("三、风险详情列表", level=1)

        if not risks:
            doc.add_paragraph("暂无风险记录")
        else:
            grouped = _group_risks(risks)
            for rule_type, items in grouped:
                label = RULE_TYPE_LABELS.get(rule_type, rule_type)
                doc.add_heading(label, level=2)

                for risk in items:
                    matched_label = MATCHED_LABELS.get(risk.matched or 0, "未知")
                    level_label = RISK_LEVEL_LABELS.get(risk.risk_level, risk.risk_level)

                    rows = [
                        ("规则名称", risk.rule_name or risk.risk_type),
                        ("严重程度", level_label),
                        ("执行结果", matched_label),
                    ]

                    if risk.matched == 2 and risk.matched_items:
                        display_items, total = _truncate_matched_items(risk.matched_items)
                        items_text = "、".join(str(i) for i in display_items)
                        if total > MAX_MATCHED_ITEMS:
                            items_text += f"（共 {total} 条命中）"
                        rows.append(("命中内容", items_text))

                    if risk.execution_result:
                        rows.append(("分析说明", risk.execution_result))

                    tbl = doc.add_table(rows=len(rows), cols=2)
                    tbl.style = "Table Grid"
                    for i, (key, val) in enumerate(rows):
                        tbl.cell(i, 0).text = key
                        tbl.cell(i, 1).text = str(val)

                    doc.add_paragraph()

        doc.save(str(report_path))


# 单例
report_generator = ReportGenerator()
