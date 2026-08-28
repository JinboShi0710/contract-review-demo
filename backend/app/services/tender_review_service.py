# -*- coding: utf-8 -*-
"""招标文件审查第一阶段。

方法借鉴 tender-review-kit（MIT）：确定性关键词撒网、带行号证据、模型归纳、
证据反查和 Excel 工作清单。关键词只作为模型的候选线索，不直接作为审核结论。
"""
import json
import re
import uuid
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.tender import TenderReviewItem, TenderReviewTask
from app.utils.llm.client import llm_client


CATEGORIES = {
    "disqualification": "否决与废标事项",
    "scoring": "评分与加分事项",
    "materials": "证明材料清单",
    "key_parameters": "重点参数与标识",
    "timeline": "时间节点",
    "contract_terms": "合同条款要点",
    "technical_requirements": "技术与功能需求",
    "acceptance_delivery": "验收、交付与运维",
}

DOCUMENT_TYPES = {
    "tender": "招标/采购文件",
    "procurement_requirement": "招标技术需求/采购需求说明书",
    "contract": "合同",
    "other": "其他文件",
}

# 精简自 tender-review-kit Community Edition 的通用信号词。
KEYWORDS = {
    "disqualification": ["否决", "废标", "无效投标", "不予受理", "不予接受", "取消投标资格", "视为撤回", "视为放弃", "拒收", "不被推荐", "不予考虑", "资格审查不合格", "实质性不响应"],
    "scoring": ["评分标准", "评分办法", "评审因素", "得分", "加分", "扣分", "技术评分", "商务评分", "价格分", "综合评分"],
    "materials": ["营业执照", "资质证书", "授权委托书", "法定代表人", "业绩证明", "合同复印件", "验收证明", "财务报表", "社保证明", "纳税证明", "承诺函", "检测报告"],
    "key_parameters": ["▲", "★", "关键参数", "实质性要求", "必须满足", "不得偏离", "正偏离", "负偏离", "技术参数"],
    "timeline": ["投标截止", "开标时间", "递交截止", "报名截止", "答疑截止", "保证金到账", "有效期", "交货期", "工期"],
    "contract_terms": ["付款条件", "付款方式", "履约保证金", "质保期", "违约责任", "验收标准", "知识产权", "保密义务", "争议解决", "合同解除", "价格调整"],
}


def _clean_lines(raw: List[Tuple[int, str]]) -> List[dict]:
    result = []
    line_no = 0
    for page, text in raw:
        for value in text.splitlines():
            value = re.sub(r"\s+", " ", value).strip()
            if not value:
                continue
            line_no += 1
            result.append({"line": line_no, "page": page, "text": value})
    return result


def _extract(path: Path) -> Tuple[List[dict], int]:
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(str(path))
        raw = [(index, page.extract_text() or "") for index, page in enumerate(reader.pages, 1)]
        return _clean_lines(raw), len(reader.pages)
    if path.suffix.lower() == ".docx":
        doc = Document(str(path))
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            chunks.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return _clean_lines([(1, "\n".join(chunks))]), 1
    raise ValueError("第一阶段仅支持 PDF 和 DOCX")


def _snippet(lines: List[dict], index: int) -> str:
    start, end = max(0, index - 1), min(len(lines), index + 2)
    return " ".join(item["text"] for item in lines[start:end])[:600]


def _scan(lines: List[dict]) -> List[dict]:
    found, seen = [], set()
    for index, row in enumerate(lines):
        for category, words in KEYWORDS.items():
            hits = [word for word in words if word in row["text"]]
            if not hits:
                continue
            quote = _snippet(lines, index)
            key = (category, quote)
            if key in seen:
                continue
            seen.add(key)
            found.append({
                "category": category,
                "title": "、".join(hits[:3]),
                "requirement": quote,
                "evidence_quote": row["text"][:400],
                "source_page": row["page"],
                "source_line": row["line"],
                "importance": "required" if category in ("disqualification", "key_parameters") else "attention",
                "action": "逐项确认响应材料，并由投标负责人复核。",
                "source": "keyword",
            })
    return found


def _parse_json_payload(text: str):
    """解析模型返回的 JSON，兼容 Markdown 代码块和前后说明文字。"""
    value = (text or "").strip()
    value = re.sub(r"^```(?:json)?\s*", "", value, flags=re.I)
    value = re.sub(r"\s*```$", "", value)
    decoder = json.JSONDecoder()
    for marker in ("{", "["):
        start = value.find(marker)
        if start < 0:
            continue
        try:
            payload, _ = decoder.raw_decode(value[start:])
            return payload
        except json.JSONDecodeError:
            continue
    raise ValueError("模型未返回有效 JSON")


def _parse_location_number(value) -> int:
    """兼容模型返回的 88、L88、行88、P1、第1页等位置格式。"""
    if isinstance(value, bool) or value is None:
        return 0
    if isinstance(value, (int, float)):
        return int(value)
    match = re.search(r"\d+", str(value))
    return int(match.group()) if match else 0


def _llm_review(lines: List[dict], keyword_items: List[dict]) -> dict:
    """单次审核受限长度文本；所有结论仍需通过原文证据反查。"""
    numbered = "\n".join(f"L{x['line']} P{x['page']} {x['text']}" for x in lines)
    candidates = [
        {"category": x["category"], "line": x["source_line"], "keyword": x["title"]}
        for x in keyword_items[:80]
    ]
    prompt = f"""你是投标方的招标文件语义审查助手。先判断文件类型，再决定是否生成投标响应清单。

文档类型只能是：tender（招标/采购文件）、procurement_requirement（项目需求或软件需求说明书）、contract（合同）、other（其他文件）。
document_type=tender 或 procurement_requirement 都属于审核范围。前者提取资格、否决、评分、材料、时间及合同条件；后者提取采购方提出的功能、性能、部署、验收、交付和运维要求。contract、other 必须返回空数组。
特别注意：采购“智能标书系统”时，文中的“系统应识别废标项、加分项、工期”等是被采购系统的功能要求，不是本项目自身的废标、加分或工期条款，应归入 technical_requirements。

审核清单只能使用以下类别：{json.dumps(CATEGORIES, ensure_ascii=False)}。
要求：
1. 每项必须是招标人对投标人的真实义务、评分条件、材料要求或合同条件；仅提及某个概念不算要求。
2. 每项必须逐字引用原文证据，page、line 必须来自输入标记。
3. 合并重复事项，不做投标/不投标结论，最多30项。
4. 关键词候选只是定位线索，必须经过上下文语义核验；错误候选不得输出。
5. source_page 和 source_line 返回纯整数。

返回纯 JSON 对象：
{{"document_type":"tender|procurement_requirement|contract|other","document_type_reason":"简短理由","findings":[{{"category":"...","title":"...","requirement":"...","evidence_quote":"原文逐字引用","source_page":1,"source_line":1,"importance":"required|attention|reference","action":"..."}}]}}

关键词候选：
{json.dumps(candidates, ensure_ascii=False)}

文件：
{numbered[:50000]}"""
    response = llm_client.chat([{"role": "user", "content": prompt}])
    payload = _parse_json_payload(response)
    # 兼容旧模型偶尔直接返回数组，但正式提示要求返回带文档类型的对象。
    if isinstance(payload, list):
        payload = {"document_type": "tender", "document_type_reason": "模型按招标文件处理", "findings": payload}
    if not isinstance(payload, dict):
        raise ValueError("模型返回结构不是 JSON 对象")
    document_type = payload.get("document_type")
    if document_type not in DOCUMENT_TYPES:
        raise ValueError("模型未返回有效文档类型")
    raw_items = payload.get("findings") or []
    if not isinstance(raw_items, list):
        raise ValueError("模型 findings 不是数组")
    if document_type not in ("tender", "procurement_requirement"):
        raw_items = []
    by_line = {x["line"]: x for x in lines}
    verified = []
    for item in raw_items:
        category = item.get("category")
        line = _parse_location_number(item.get("source_line"))
        evidence = str(item.get("evidence_quote") or "").strip()
        source = by_line.get(line)
        if category not in CATEGORIES or not source or not evidence:
            continue
        # 证据反查：模型引用必须能在原行或相邻上下文中找到明显片段。
        context = " ".join(by_line[n]["text"] for n in range(max(1, line - 1), line + 2) if n in by_line)
        compact_source = re.sub(r"\s+", "", context)
        compact_evidence = re.sub(r"\s+", "", evidence)
        if compact_evidence[:10] not in compact_source and compact_source[:10] not in compact_evidence:
            continue
        verified.append({
            "category": category,
            "title": str(item.get("title") or CATEGORIES[category])[:255],
            "requirement": str(item.get("requirement") or evidence),
            "evidence_quote": evidence[:600],
            "source_page": source["page"],
            "source_line": line,
            "importance": item.get("importance") if item.get("importance") in ("required", "attention", "reference") else "attention",
            "action": str(item.get("action") or "请人工确认响应方式。"),
            "source": "llm",
        })
    return {
        "document_type": document_type,
        "document_type_reason": str(payload.get("document_type_reason") or "").strip(),
        "findings": verified,
    }


class TenderReviewService:
    def upload(self, file, db: Session, created_by: str) -> dict:
        ext = Path(file.filename).suffix.lower()
        if ext not in (".pdf", ".docx"):
            raise ValueError("第一阶段仅支持 PDF、DOCX")
        content = file.file.read()
        if not content or len(content) > 50 * 1024 * 1024:
            raise ValueError("文件为空或超过 50MB")
        task_id = str(uuid.uuid4())
        folder = settings.UPLOAD_DIR / "tenders"
        folder.mkdir(parents=True, exist_ok=True)
        path = folder / f"{task_id}{ext}"
        path.write_bytes(content)
        task = TenderReviewTask(id=task_id, file_name=file.filename, file_path=str(path), file_type=ext[1:], file_size=len(content), status="created", created_by=created_by)
        db.add(task)
        db.commit()
        return {"task_id": task_id, "status": "created", "file_name": file.filename}

    def process(self, task_id: str, db: Session) -> dict:
        task = db.query(TenderReviewTask).filter(TenderReviewTask.id == task_id, TenderReviewTask.is_deleted == 0).first()
        if not task:
            raise ValueError("审查任务不存在")
        task.status, task.error_message = "processing", None
        db.commit()
        try:
            lines, pages = _extract(Path(task.file_path))
            if sum(len(x["text"]) for x in lines) < 80:
                raise ValueError("文件可提取文字过少，可能是扫描件；第一阶段请使用可复制文字的 PDF 或 DOCX")
            keyword_items = _scan(lines)
            try:
                llm_result = _llm_review(lines, keyword_items)
            except Exception as exc:
                raise RuntimeError(f"大模型语义审核失败，未生成审核结论：{exc}") from exc
            document_type = llm_result["document_type"]
            merged = llm_result["findings"]
            db.query(TenderReviewItem).filter(TenderReviewItem.task_id == task_id).delete()
            for item in merged:
                db.add(TenderReviewItem(task_id=task_id, **item))
            task.title = next((x["text"] for x in lines[:8] if len(x["text"]) > 4), task.file_name)
            task.extracted_text = "\n".join(f"L{x['line']} P{x['page']} {x['text']}" for x in lines)
            task.page_count, task.line_count, task.status = pages, len(lines), "completed"
            type_label = DOCUMENT_TYPES[document_type]
            type_reason = llm_result["document_type_reason"]
            if document_type in ("tender", "procurement_requirement"):
                task.summary = f"AI识别为{type_label}，已形成 {len(merged)} 条有原文证据的投标响应事项。{type_reason}"
            else:
                task.summary = f"AI识别为{type_label}，不属于招标审核适用范围，未生成投标响应清单。{type_reason}"
            db.commit()
            return self.detail(task_id, db)
        except Exception as exc:
            task.status, task.error_message = "failed", str(exc)
            db.commit()
            raise

    def _task(self, task_id: str, db: Session):
        return db.query(TenderReviewTask).filter(TenderReviewTask.id == task_id, TenderReviewTask.is_deleted == 0).first()

    def detail(self, task_id: str, db: Session) -> dict:
        task = self._task(task_id, db)
        if not task:
            raise ValueError("审查任务不存在")
        items = db.query(TenderReviewItem).filter(TenderReviewItem.task_id == task_id).order_by(TenderReviewItem.category, TenderReviewItem.source_page, TenderReviewItem.source_line).all()
        return {"id": task.id, "file_name": task.file_name, "title": task.title, "status": task.status, "page_count": task.page_count, "line_count": task.line_count, "summary": task.summary, "error_message": task.error_message, "created_at": task.created_at.isoformat(), "categories": CATEGORIES, "items": [{"id": x.id, "category": x.category, "title": x.title, "requirement": x.requirement, "evidence_quote": x.evidence_quote, "source_page": x.source_page, "source_line": x.source_line, "importance": x.importance, "action": x.action, "source": x.source} for x in items]}

    def list(self, db: Session, user_id: str, role: str) -> dict:
        query = db.query(TenderReviewTask).filter(TenderReviewTask.is_deleted == 0)
        if role != "admin":
            query = query.filter(TenderReviewTask.created_by == user_id)
        tasks = query.order_by(TenderReviewTask.created_at.desc()).all()
        return {"items": [{"id": x.id, "file_name": x.file_name, "title": x.title, "status": x.status, "summary": x.summary, "created_at": x.created_at.isoformat()} for x in tasks], "total": len(tasks)}

    def export(self, task_id: str, db: Session) -> Path:
        data = self.detail(task_id, db)
        folder = settings.EXPORT_DIR / "tenders"
        folder.mkdir(parents=True, exist_ok=True)
        path = folder / f"tender_review_{task_id[:8]}.xlsx"
        wb = Workbook()
        wb.remove(wb.active)
        colors = {"disqualification": "C00000", "scoring": "00B050", "materials": "7030A0", "key_parameters": "ED7D31", "timeline": "4472C4", "contract_terms": "BF8F00", "technical_requirements": "2F75B5", "acceptance_delivery": "548235"}
        for key, label in CATEGORIES.items():
            ws = wb.create_sheet(label[:31])
            headers = ["事项", "要求/说明", "原文证据", "页码", "行号", "重要程度", "建议动作", "责任人", "完成情况", "备注"]
            ws.append(headers)
            for cell in ws[1]:
                cell.fill, cell.font = PatternFill("solid", fgColor=colors[key]), Font(color="FFFFFF", bold=True)
            for item in data["items"]:
                if item["category"] == key:
                    ws.append([item["title"], item["requirement"], item["evidence_quote"], item["source_page"], item["source_line"], item["importance"], item["action"], "", "", ""])
            ws.freeze_panes, ws.auto_filter.ref = "A2", ws.dimensions
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
            for col in "ABCDEFGHIJ":
                ws.column_dimensions[col].width = 16 if col not in "BCG" else 42
        wb.save(path)
        return path


tender_review_service = TenderReviewService()
